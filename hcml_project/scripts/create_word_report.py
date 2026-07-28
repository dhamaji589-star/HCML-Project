"""Create a polished Word report using python-docx."""

from __future__ import annotations

import csv
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


OUTPUT = Path("hcml_project/docs/HCML_MAD22_Report_Draft.docx")
SUMMARY_CSV = Path("results/summary_method_results.csv")
SUCCESS_FIGURE = Path("results/success_rate_by_method_horizontal.png")
QUALITATIVE_FIGURE = Path("results/qualitative_best_examples.png")


def set_columns(section, count: int, space_twips: int = 420) -> None:
    sect_pr = section._sectPr
    cols = sect_pr.xpath("./w:cols")
    cols = cols[0] if cols else OxmlElement("w:cols")
    cols.set(qn("w:num"), str(count))
    cols.set(qn("w:space"), str(space_twips))
    if not sect_pr.xpath("./w:cols"):
        sect_pr.append(cols)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(8.2)


def add_para(doc: Document, text: str, style: str | None = None, bold_start: str | None = None):
    paragraph = doc.add_paragraph(style=style)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.03
    if bold_start and text.startswith(bold_start):
        first = paragraph.add_run(bold_start)
        first.bold = True
        rest = paragraph.add_run(text[len(bold_start) :])
        for run in [first, rest]:
            run.font.name = "Times New Roman"
            run.font.size = Pt(9.5)
    else:
        run = paragraph.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(9.5)
    return paragraph


def add_heading(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(7)
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(text)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(10.2)


def add_caption(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    run.italic = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(8.2)


def success_to_percent(value: str) -> float:
    successes, trials = value.split("/")
    return 100.0 * int(successes) / int(trials)


def read_summary() -> list[dict[str, str]]:
    with SUMMARY_CSV.open("r", newline="", encoding="utf-8") as csv_file:
        return list(csv.DictReader(csv_file))


def create_horizontal_success_plot() -> None:
    rows = read_summary()
    methods = [row["method"] for row in rows]
    neg = [success_to_percent(row["negfacediff_success"]) for row in rows]
    adapt = [success_to_percent(row["adaptdiff_success"]) for row in rows]

    width, height = 1100, 620
    left, right, top, bottom = 165, 70, 55, 55
    row_gap = 92
    bar_height = 24
    max_value = 100
    plot_width = width - left - right
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    try:
        title_font = ImageFont.truetype("arial.ttf", 30)
        font = ImageFont.truetype("arial.ttf", 23)
        small_font = ImageFont.truetype("arial.ttf", 19)
    except OSError:
        title_font = ImageFont.load_default()
        font = ImageFont.load_default()
        small_font = ImageFont.load_default()

    draw.text((left, 18), "Hidden-identity recovery success rate", fill="black", font=title_font)
    draw.rectangle((left + 565, 24, left + 590, 49), fill="#4C78A8")
    draw.text((left + 600, 22), "NegFaceDiff", fill="black", font=small_font)
    draw.rectangle((left + 745, 24, left + 770, 49), fill="#F58518")
    draw.text((left + 780, 22), "AdaptDiff", fill="black", font=small_font)

    axis_top = top + 35
    for tick in range(0, 101, 20):
        x = left + tick / max_value * plot_width
        draw.line((x, axis_top, x, height - bottom), fill="#DDDDDD", width=1)
        draw.text((x - 14, height - bottom + 10), str(tick), fill="black", font=small_font)
    draw.text((left + plot_width / 2 - 65, height - 25), "Success (%)", fill="black", font=small_font)

    for index, method in enumerate(methods):
        y = axis_top + index * row_gap + 20
        draw.text((18, y + 11), method, fill="black", font=font)
        for offset, value, color in [(0, neg[index], "#4C78A8"), (bar_height + 8, adapt[index], "#F58518")]:
            bar_y = y + offset
            bar_width = value / max_value * plot_width
            draw.rectangle((left, bar_y, left + bar_width, bar_y + bar_height), fill=color)
            draw.text((left + bar_width + 8, bar_y - 1), f"{value:.1f}%", fill="black", font=small_font)

    SUCCESS_FIGURE.parent.mkdir(parents=True, exist_ok=True)
    image.save(SUCCESS_FIGURE)


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(9.5)

    for style_name in ["Title", "Heading 1", "Heading 2"]:
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style.font.color.rgb = RGBColor(0, 0, 0)


def add_title_and_abstract(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(2)
    run = title.add_run("Hidden Identity Recovery from Face Morphing Attacks using NegFaceDiff and AdaptDiff")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(15)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(8)
    run = meta.add_run("Vishu | HCML MAD22 Project")
    run.italic = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(9.5)

    add_heading(doc, "Abstract")
    add_para(
        doc,
        "Face morphing attacks combine identity information from two subjects into one facial image. "
        "This project investigates hidden-identity recovery: given a morph image and one known contributor, "
        "the objective is to generate a sample that moves toward the other contributor. The implemented "
        "pipeline uses pretrained NegFaceDiff/AdaptDiff components, the DM_CASIA latent diffusion model, a "
        "pretrained latent autoencoder, and ElasticFaceArc identity embeddings. Five MAD22 morphing subsets "
        "were evaluated independently: OpenCV, FaceMorpher, MIPGAN-I, MIPGAN-II, and WebMorph. AdaptDiff "
        "achieved higher recovery success and higher mean hidden-minus-known identity margins for every "
        "morphing method. The generated samples are visually imperfect, so the conclusion is framed as "
        "identity recovery in embedding space rather than photorealistic reconstruction.",
    )


def add_results_table(doc: Document) -> None:
    rows = [
        ["Morphing method", "NegFaceDiff success", "NegFaceDiff margin", "AdaptDiff success", "AdaptDiff margin"],
        ["OpenCV", "60.0%", "0.069954", "62.5%", "0.090291"],
        ["FaceMorpher", "75.0%", "0.045358", "82.5%", "0.059136"],
        ["MIPGAN-I", "80.0%", "0.063586", "88.8%", "0.084853"],
        ["MIPGAN-II", "72.5%", "0.060798", "80.0%", "0.078214"],
        ["WebMorph", "71.2%", "0.063933", "75.0%", "0.087545"],
    ]
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    table.autofit = True
    for row_index, row in enumerate(rows):
        for col_index, text in enumerate(row):
            set_cell_text(table.cell(row_index, col_index), text, bold=row_index == 0)
            if row_index == 0:
                set_cell_shading(table.cell(row_index, col_index), "EDEDED")
    add_caption(
        doc,
        "Table 1. Success is reported as a percentage. The margin is the mean cosine(generated, hidden) "
        "minus cosine(generated, known).",
    )


def add_references(doc: Document) -> None:
    refs = [
        "[1] E. Caldeira, T. Chettaoui, N. Damer, and F. Boutros, AdaptDiff: Adaptive Guidance in Diffusion Models for Diverse and Identity-Consistent Face Synthesis, 2026. Code: https://github.com/EduardaCaldeira/NegFaceDiff/",
        "[2] M. Huber et al., SYN-MAD 2022: Competition on Face Morphing Attack Detection Based on Privacy-aware Synthetic Training Data, IJCB, 2022. Code: https://github.com/marcohuber/SYN-MAD-2022",
        "[3] F. Boutros, N. Damer, F. Kirchbuchner, and A. Kuijper, ElasticFace: Elastic Margin Loss for Deep Face Recognition, CVPR Workshops, 2022.",
        "[4] J. Deng, J. Guo, N. Xue, and S. Zafeiriou, ArcFace: Additive Angular Margin Loss for Deep Face Recognition, CVPR, 2019.",
        "[5] J. Ho, A. Jain, and P. Abbeel, Denoising Diffusion Probabilistic Models, NeurIPS, 2020.",
        "[6] J. Song, C. Meng, and S. Ermon, Denoising Diffusion Implicit Models, ICLR, 2021.",
        "[7] R. Rombach, A. Blattmann, D. Lorenz, P. Esser, and B. Ommer, High-Resolution Image Synthesis with Latent Diffusion Models, CVPR, 2022.",
    ]
    add_heading(doc, "References")
    for ref in refs:
        paragraph = add_para(doc, ref)
        paragraph.paragraph_format.left_indent = Cm(0.35)
        paragraph.paragraph_format.first_line_indent = Cm(-0.35)


def build_report() -> None:
    create_horizontal_success_plot()
    doc = Document()
    configure_styles(doc)

    section = doc.sections[0]
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.45)
    section.right_margin = Cm(1.45)

    add_title_and_abstract(doc)

    body_section = doc.add_section(WD_SECTION.CONTINUOUS)
    set_columns(body_section, 2)

    add_heading(doc, "1. Introduction")
    add_para(
        doc,
        "Face morphing is a biometric attack in which two identities are blended into a single face image. "
        "A successful morph can contain enough information from both contributors to match against either "
        "person in a face-recognition system. In this work, the goal is not to classify an image as bona fide "
        "or morphed. Instead, the project studies whether the hidden contributor can be recovered when the "
        "morph image and the other contributor are already known.",
    )
    add_para(
        doc,
        "The task is directional. For a morph created from identities A and B, one trial uses A as the known "
        "identity and treats B as the hidden target, while the reverse trial uses B as known and A as hidden. "
        "This formulation makes the guidance objective clear: use the mixed identity information contained in "
        "the morph, but push the generated sample away from the known contributor.",
    )

    add_heading(doc, "2. Method")
    add_para(
        doc,
        "The complete pipeline is built from pretrained components and performs inference only. Dataset "
        "filenames are first converted into trial metadata containing the morph image, the known identity, and "
        "the hidden identity. ElasticFaceArc with an iresnet100 backbone then extracts 512-dimensional identity "
        "embeddings. The morph embedding is used as the positive context, the known identity embedding as the "
        "negative context, and the hidden identity embedding only for evaluation.",
    )
    add_para(
        doc,
        "Image generation is performed in latent space. The morph image is encoded by the pretrained first-stage "
        "autoencoder. Forward diffusion noise is added using the 1000-step schedule associated with the provided "
        "DM_CASIA model. DDIM sampling then performs 200 reverse denoising steps, and the final latent is decoded "
        "back into an image. This follows the practical shortcut recommended for sampling while keeping the "
        "starting noise consistent with the model setup.",
    )
    add_para(
        doc,
        "Two guidance strategies are compared. NegFaceDiff uses fixed negative identity guidance with adapt=false "
        "and weight=0.5. AdaptDiff uses adaptive negative guidance with adapt=true and weight=1.0. Both methods "
        "try to retain identity information from the morph while suppressing the known contributor, but AdaptDiff "
        "changes the strength of negative guidance during the denoising trajectory.",
    )
    add_para(
        doc,
        "The implementation was organized as a sequence of small scripts rather than one large notebook cell. "
        "This made the workflow easier to verify: one script builds trial metadata, one builds the unique image "
        "manifest, one aligns faces, one extracts ElasticFaceArc embeddings, one prepares paired contexts for "
        "sampling, and one runs the diffusion model. This structure was helpful because intermediate CSV files "
        "could be inspected before using GPU time for generation.",
    )

    add_heading(doc, "3. Experimental Protocol")
    add_para(
        doc,
        "The evaluated MAD22 morphing subsets are OpenCV, FaceMorpher, MIPGAN-I, MIPGAN-II, and WebMorph. Each "
        "subset is evaluated independently because morphing methods can distribute contributor identities in "
        "different ways. Reporting them separately prevents one method from hiding the behavior of another.",
    )
    add_para(
        doc,
        "For each directed trial, both NegFaceDiff and AdaptDiff generate one candidate recovery image. The "
        "generated image is embedded with ElasticFaceArc and compared with the known and hidden identity "
        "embeddings using cosine similarity. A trial is successful when cosine(generated, hidden) is greater "
        "than cosine(generated, known). The margin is the difference between these two similarities. A larger "
        "positive margin indicates a clearer movement toward the hidden identity in embedding space.",
    )
    add_para(
        doc,
        "The success metric is intentionally relative. It does not require the generated image to perfectly match "
        "the hidden identity; it asks whether the image is closer to the hidden contributor than to the known "
        "contributor. This is suitable for this project because the objective is identity disentanglement from a "
        "morph, not exact image reconstruction. The margin complements the success rate by showing how confident "
        "this relative movement is on average.",
    )

    add_heading(doc, "4. Quantitative Results")
    doc.add_picture(str(SUCCESS_FIGURE), width=Inches(3.05))
    add_caption(doc, "Figure 1. Hidden-identity recovery success rate for each morphing method.")
    add_results_table(doc)
    add_para(
        doc,
        "AdaptDiff improves over NegFaceDiff for all five morphing methods. The improvement is visible in the "
        "success percentage and in the mean margin. This is important because the margin measures more than a "
        "binary success decision: it shows how strongly the generated embedding is separated from the known "
        "identity and shifted toward the hidden contributor.",
    )
    add_para(
        doc,
        "MIPGAN-I is the easiest subset in this evaluation, while OpenCV is the hardest. This variation supports "
        "the decision to keep results separated by morphing method. Different morphing algorithms appear to "
        "leave different amounts or forms of recoverable identity information in the generated morph image.",
    )
    add_para(
        doc,
        "Another useful observation is that AdaptDiff does not improve only one isolated subset. Its advantage is "
        "consistent across classical morphs and GAN-based morphs. This consistency makes the result more reliable "
        "than a single high score on one subset, because it suggests that adaptive guidance is useful under several "
        "different morph generation procedures.",
    )

    add_heading(doc, "5. Qualitative Results")
    add_para(
        doc,
        "The qualitative examples were selected from successful AdaptDiff cases with strong positive margins. "
        "They are useful because they show the complete recovery setup: known identity, input morph, hidden "
        "identity, and both generated outputs.",
    )

    figure_section = doc.add_section(WD_SECTION.CONTINUOUS)
    set_columns(figure_section, 1)
    doc.add_picture(str(QUALITATIVE_FIGURE), width=Inches(6.25))
    add_caption(
        doc,
        "Figure 2. Qualitative examples selected by positive AdaptDiff margin. The generated faces show an "
        "identity signal but remain visually blurry.",
    )

    body_section_2 = doc.add_section(WD_SECTION.CONTINUOUS)
    set_columns(body_section_2, 2)
    add_para(
        doc,
        "The visual limitation is clear. The generated samples are often blurry, color-shifted, and less realistic "
        "than the original MAD22 images. Therefore, the method should not be described as producing clean "
        "photo-quality reconstructions. A more accurate interpretation is that the generated samples contain "
        "identity cues that ElasticFaceArc places closer to the hidden contributor than to the known contributor.",
    )
    add_para(
        doc,
        "This gap between visual quality and embedding behavior is important in face analysis. Human observers "
        "mainly judge sharpness, lighting, expression, and realism, while the evaluation model compares numerical "
        "identity features. A blurry output can therefore still be useful for the metric if it contains enough "
        "identity-related structure. For this reason, the qualitative and quantitative results should be read "
        "together rather than treated as interchangeable evidence.",
    )

    add_heading(doc, "6. Discussion")
    add_para(
        doc,
        "The consistent advantage of AdaptDiff suggests that adaptive negative guidance is better suited to this "
        "recovery task than fixed negative guidance. Diffusion sampling changes character over time: early steps "
        "shape coarse structure, while later steps refine identity and appearance. Fixed negative guidance can "
        "constrain the process uniformly, whereas adaptive guidance can allow broader exploration early and "
        "stronger identity separation later.",
    )
    add_para(
        doc,
        "The main limitations are also important. First, no component was fine-tuned for this exact recovery task. "
        "Second, ElasticFaceArc is used both for identity conditioning and for evaluation, so an independent "
        "face-recognition model would be useful as a future check. Third, the generated images do not yet have "
        "strong visual fidelity. These limitations mean the project should be framed as evidence of hidden-identity "
        "recovery in embedding space, not as a final visual reconstruction system.",
    )
    add_para(
        doc,
        "A natural next step would be to repeat the evaluation with an additional face-recognition model that was "
        "not used for conditioning. This would test whether the recovered identity signal transfers beyond "
        "ElasticFaceArc. Another extension would be to study different noise timesteps, guidance weights, and "
        "sampling schedules more systematically. In the current project, the settings were kept close to the "
        "supervisor's instructions so that the comparison between NegFaceDiff and AdaptDiff remained controlled.",
    )

    add_heading(doc, "7. Conclusion")
    add_para(
        doc,
        "This project implemented an end-to-end hidden-identity recovery pipeline for MAD22 morphing attacks. The "
        "pipeline prepares directed trials, extracts ElasticFaceArc identity contexts, samples with the pretrained "
        "latent diffusion model, decodes generated images, and evaluates them using cosine similarity in identity "
        "embedding space. Across all evaluated morphing methods, AdaptDiff performs better than NegFaceDiff in "
        "both success percentage and mean identity margin. The results support adaptive negative guidance as the "
        "stronger strategy for this experimental setting, while also showing that visual realism remains an open "
        "challenge.",
    )

    refs_section = doc.add_section(WD_SECTION.NEW_PAGE)
    set_columns(refs_section, 1)
    add_references(doc)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"Word report written: {OUTPUT}")
    print(f"Horizontal success figure written: {SUCCESS_FIGURE}")


if __name__ == "__main__":
    build_report()
